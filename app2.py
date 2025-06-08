import streamlit as st 
#from text_extraction import extract_text_from_file
from generate_mom import generate_minutes_of_meeting, extract_text_from_image
from formatting import generate_mom_html
from PIL import Image
import pandas as pd
import io
import tempfile
from xhtml2pdf import pisa
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io

# App Title
st.header("📋 :blue[AI] M.O.M Generator (Multi-Format)")
st.subheader("🔰 Tips for Using this Application: ")
notes = f'''
* **Upload Your Minutes:** The first step is to upload the hand written Minutes/Call Notes as image.
* **Click "Generate MoM":** The app will extract, summarize, and generate all discussion points and action items in a structured way.'''
st.write(notes)

# Upload section
uploaded_file = st.file_uploader("Upload meeting notes file here...", 
                                 type=["jpg", "png", "pdf", "docx", "txt"])

if uploaded_file:
    # Display uploaded image if image
    file_type = uploaded_file.type
    if file_type.startswith("image/"):
        image = Image.open(uploaded_file)
        st.image(image, caption="Uploaded Handwritten Notes", use_container_width=True)
        
        # Save temp image file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp:
            image.save(tmp.name)
            img_path = tmp.name

        with st.spinner("🔍 Extracting content from handwritten image..."):
            raw_text = extract_text_from_image(img_path)
        st.subheader("📄 :orange[Extracted Text]")
        st.text_area("Raw Text", raw_text, height=300)
        
def generate_word_file(response_text: str) -> io.BytesIO:
    
    """
     Converts the AI response text (which includes a table and a summary section)
    into a formatted Word document.
    """
    doc = Document()

    # Title
    title = doc.add_heading("Minutes of Meeting", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    lines = response_text.strip().split('\n')
    table_started = False
    summary_started = False

    table = None

    for line in lines:
        line = line.strip()

        if line.lower().startswith("date:"):
            doc.add_paragraph(line)

        elif "|" in line and not summary_started:
            parts = [cell.strip() for cell in line.strip('|').split('|')]
            if not table_started:
                table = doc.add_table(rows=1, cols=len(parts))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, part in enumerate(parts):
                    hdr_cells[i].text = part
                table_started = True
            else:
                row_cells = table.add_row().cells
                for i, part in enumerate(parts):
                    row_cells[i].text = part

        elif "Summary & Key Action Items" in line:
            doc.add_paragraph()
            doc.add_heading("Summary & Key Action Items", level=1)
            summary_started = True

        elif summary_started:
            if line:
                para = doc.add_paragraph(line)
                para.style.font.size = Pt(11)

    # Save to buffer
    word_io = io.BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io

if st.button("🧠 Generate MoM using AI"):
    with st.spinner("⏳ AI is Working..."):
        formatted_mom = generate_minutes_of_meeting(raw_text)
            #st.subheader("✅ Structured Minutes of Meeting")
            #st.code(formatted_mom)
            #st.write(formatted_mom)
            #st.write(json.dumps(formatted_mom, indent=2))

            # Generate HTML
            #html = generate_mom_html(formatted_mom)

            # Convert to PDF
            #pdf_file = io.BytesIO()
            #pisa.CreatePDF(io.StringIO(html), dest=pdf_file)
            #pdf_file.seek(0)

            # Download PDF
        docx_file = generate_word_file(formatted_mom)
        st.download_button(label="📄📥 Download Word File (.docx)",data=docx_file,
                file_name="Minutes_of_Meeting.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")